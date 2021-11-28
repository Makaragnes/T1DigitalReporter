import cv2
import docx
import pyrebase
import json
import pytesseract
from docx import Document

config = {
  "apiKey": "AIzaSyCSP1Ps1Gn2l0QSeFIwCNViLy52wcJab20",
  "authDomain": "t1reporter.firebaseapp.com",
  "databaseURL": "https://t1reporter-default-rtdb.europe-west1.firebasedatabase.app",
  "projectId": "t1reporter",
  "storageBucket": "t1reporter.appspot.com",
  "messagingSenderId": "960754275384",
  "appId": "1:960754275384:web:d6e506b30fd913087842cd",
  "measurementId": "G-E5HR173BX1"
}

''' "apiKey": "AIzaSyARwzPEfYY_YvVCS8RZfj8N87OtFIFoZCM",
  "authDomain": "ahih-bef37.firebaseapp.com",
  "databaseURL": "https://ahih-bef37-default-rtdb.firebaseio.com",
  "projectId": "ahih-bef37",
  "storageBucket": "ahih-bef37.appspot.com",
  "messagingSenderId": "535684827832",
  "appId": "1:535684827832:web:cbd3d56d8c83195335c24d",
  "measurementId": "G-SCNT5QCMRB",
  "serviceAccount": 'cert.json'
'''

'''firebase = pyrebase.initialize_app(config)

db = firebase.database()
data = {"fam": "Aliuy Alievich"}
db.child("Users").child("fweguowh").child("Users").push(data)
ahih=db.child("Users").get()
for user in ahih.each():
  print(user.val())
print("suc")'''

# Opening JSON file
'''
f = open('google-services.json')
# returns JSON object as
# a dictionary
data = json.load(f)
# Iterating through the json
# list
for i in data['client']:
  for j in i['client_info']:
    print(j)
# Closing file
f.close()
'''
firebase = pyrebase.initialize_app(config)
storage = firebase.storage()
db = firebase.database()

storage.child().download("14.jpg")

'''storage.child("users/Rzhevskiy/1.jpg").put("1.jpg")
#storage.child("users/Martin/2.jpg").put("2.jpg")
storage.child("users/Rzhevskiy/3.jpg").put("3.jpg")
#storage.child("users/Martin/4.png").put("4.png")
storage.child("users/Rzhevskiy/5.png").put("5.png")
#storage.child("users/Martin/6.png").put("6.png")
'''

#storage.child("users/Martin/")
#files = storage.list_files()

#for file in files:
#    print(storage.child(file.name).get_url(None))

img = cv2.imread('5.png')
img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
print(pytesseract.get_languages)
config = r'--oem 3 --psm 6'
data = {"Text": pytesseract.image_to_string(img, lang='rus', config=config)}
db.child("Users").child("textFromImages").push(data)

doc = Document()
doc.add_paragraph('some more text ... ')
doc.add_paragraph('sqwswdc ome fw more text ... ')
doc.save("test.docx")


''' 
auth = firebase.auth()

runtimeDatabase = firebase.storage()

email = input("email")

password = input("password")

user = auth.create_user_with_email_and_password(email, password)
'''


